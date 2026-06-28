from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from backend.app.config import ensure_directories, settings
from backend.app.services.processing import MaterialNotFoundError
from backend.app.services.report_generator import ReportGenerationError, normalize_report_type
from backend.app.services.storage import get_material, load_materials, save_materials


class ExportError(ValueError):
    pass


REPORT_EXPORT_FILENAMES = {
    "short": {
        "pdf": "short_report.pdf",
        "docx": "short_report.docx",
    },
    "full_clean": {
        "pdf": "full_clean_notes.pdf",
        "docx": "full_clean_notes.docx",
    },
}

UNICODE_FONT_CANDIDATES = [
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
    Path("/Library/Fonts/Arial.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
]


def _save_material_metadata(metadata: dict[str, Any]) -> None:
    materials = load_materials()
    materials[metadata["material_id"]] = metadata
    save_materials(materials)


def _load_material(material_id: str) -> dict[str, Any]:
    metadata = get_material(material_id)
    if metadata is None:
        raise MaterialNotFoundError("Material not found")
    return metadata


def _report_key(report_type: str) -> str:
    try:
        return normalize_report_type(report_type)
    except ReportGenerationError as exc:
        raise ExportError("Invalid report_type") from exc


def get_export_paths(material_id: str, report_type: str) -> dict:
    report_key = _report_key(report_type)
    reports_dir = Path(settings.REPORTS_DIR) / material_id
    return {
        "report_key": report_key,
        "reports_dir": reports_dir,
        "pdf_path": reports_dir / REPORT_EXPORT_FILENAMES[report_key]["pdf"],
        "docx_path": reports_dir / REPORT_EXPORT_FILENAMES[report_key]["docx"],
    }


def get_report_markdown_path(material_metadata: dict, report_type: str) -> Path:
    report_key = _report_key(report_type)
    report = (material_metadata.get("reports") or {}).get(report_key)
    if not report or not report.get("path"):
        raise ExportError("Markdown report does not exist. Generate the report first.")

    markdown_path = Path(report["path"])
    if not markdown_path.exists():
        raise ExportError("Markdown report does not exist. Generate the report first.")
    return markdown_path


def _update_export_metadata(
    metadata: dict[str, Any],
    report_key: str,
    format_name: str,
    export_path: Path,
) -> None:
    reports = metadata.get("reports") or {}
    report = reports.get(report_key) or {}
    report[f"{format_name}_path"] = str(export_path)
    reports[report_key] = report
    metadata["reports"] = reports
    _save_material_metadata(metadata)


def _write_docx(markdown: str, output_path: Path) -> None:
    document = Document()

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)

    document.save(output_path)


def _register_pdf_font() -> str:
    for font_path in UNICODE_FONT_CANDIDATES:
        if font_path.exists():
            font_name = "ReportUnicodeFont"
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
            return font_name

    # Helvetica is always available in ReportLab, but Cyrillic support may be limited.
    return "Helvetica"


def _build_pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    sample_styles = getSampleStyleSheet()
    return {
        "normal": ParagraphStyle(
            "ReportNormal",
            parent=sample_styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=15,
            alignment=TA_LEFT,
            spaceAfter=4,
        ),
        "h1": ParagraphStyle(
            "ReportHeading1",
            parent=sample_styles["Heading1"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            spaceBefore=8,
            spaceAfter=10,
        ),
        "h2": ParagraphStyle(
            "ReportHeading2",
            parent=sample_styles["Heading2"],
            fontName=font_name,
            fontSize=14,
            leading=19,
            spaceBefore=8,
            spaceAfter=7,
        ),
        "h3": ParagraphStyle(
            "ReportHeading3",
            parent=sample_styles["Heading3"],
            fontName=font_name,
            fontSize=12,
            leading=16,
            spaceBefore=6,
            spaceAfter=5,
        ),
    }


def _flush_pdf_bullets(
    story: list,
    bullet_items: list[ListItem],
) -> None:
    if not bullet_items:
        return

    story.append(
        ListFlowable(
            bullet_items,
            bulletType="bullet",
            leftIndent=12,
            bulletFontSize=8,
        )
    )
    story.append(Spacer(1, 4))
    bullet_items.clear()


def _write_pdf(markdown: str, output_path: Path) -> None:
    font_name = _register_pdf_font()
    styles = _build_pdf_styles(font_name)
    story = []
    bullet_items: list[ListItem] = []

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            _flush_pdf_bullets(story, bullet_items)
            story.append(Spacer(1, 6))
        elif stripped.startswith("# "):
            _flush_pdf_bullets(story, bullet_items)
            story.append(Paragraph(escape(stripped[2:].strip()), styles["h1"]))
        elif stripped.startswith("## "):
            _flush_pdf_bullets(story, bullet_items)
            story.append(Paragraph(escape(stripped[3:].strip()), styles["h2"]))
        elif stripped.startswith("### "):
            _flush_pdf_bullets(story, bullet_items)
            story.append(Paragraph(escape(stripped[4:].strip()), styles["h3"]))
        elif stripped.startswith("- "):
            bullet_items.append(
                ListItem(Paragraph(escape(stripped[2:].strip()), styles["normal"]))
            )
        else:
            _flush_pdf_bullets(story, bullet_items)
            story.append(Paragraph(escape(stripped), styles["normal"]))

    _flush_pdf_bullets(story, bullet_items)

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    document.build(story)


def export_report_to_pdf(material_id: str, report_type: str) -> Path:
    metadata = _load_material(material_id)
    markdown_path = get_report_markdown_path(metadata, report_type)
    paths = get_export_paths(material_id, report_type)

    ensure_directories()
    paths["reports_dir"].mkdir(parents=True, exist_ok=True)
    markdown = markdown_path.read_text(encoding="utf-8")
    _write_pdf(markdown, paths["pdf_path"])
    _update_export_metadata(metadata, paths["report_key"], "pdf", paths["pdf_path"])
    return paths["pdf_path"]


def export_report_to_docx(material_id: str, report_type: str) -> Path:
    metadata = _load_material(material_id)
    markdown_path = get_report_markdown_path(metadata, report_type)
    paths = get_export_paths(material_id, report_type)

    ensure_directories()
    paths["reports_dir"].mkdir(parents=True, exist_ok=True)
    markdown = markdown_path.read_text(encoding="utf-8")
    _write_docx(markdown, paths["docx_path"])
    _update_export_metadata(metadata, paths["report_key"], "docx", paths["docx_path"])
    return paths["docx_path"]
